import warnings

import subprocess

import time

import database
from docx import Document

from flask import Flask, jsonify, request, abort

warnings.filterwarnings("ignore")

app = Flask(__name__)

items = []


@app.route('/')
def index():
    return "Hello, World!"


@app.route('/api/v1.0/matiere', methods=['GET'])
def get_matieres():
    database.getmatieres()
    result = database.resultsExportMatieres
    print (result)

    return jsonify({'item': result}), 201


@app.route('/api/v1.0/matiere', methods=['POST'])
def create_matiere():
    database.creatematiere(request.json)

    return jsonify({'item': 'matiere cree'}), 201


@app.route('/api/v1.0/etudiant', methods=['GET'])
def get_etudiants():
    database.getetudiants()
    result = database.resultsExportEtudiants
    print (result)
    return jsonify({'item': result}), 201


@app.route('/api/v1.0/etudiant', methods=['POST'])
def create_etudiant():
    database.createetudiant(request.json)

    return jsonify({'item': 'etudiant cree'}), 201


@app.route('/api/v1.0/note', methods=['GET'])
def get_notes():
    database.getnotes()
    result = database.resultsExportNotes
    print (result)
    return jsonify({'item': result}), 201


@app.route('/api/v1.0/note', methods=['POST'])
def create_note():
    database.createnote(request.json)

    return jsonify({'item': 'note cree'}), 201


@app.route('/api/v1.0/noteetu/', methods=['GET'])
def get_notes_etu():
    #if not request.json or not 'id_etu' in request.args.get('username'):
    #    abort(400)
    database.getnotesetu(int(request.args.get('id_etu')))
    result = database.resultsExportNotesEtu
    document = Document()

    document.add_heading('Bulletin MASTER 2 INFO', 0)

    document.add_paragraph('')
    document.add_paragraph('Etudiant')
    document.add_paragraph('Matricule: ' + result[0]['mat'])
    document.add_paragraph('Nom: ' + result[0]['nom'])
    document.add_paragraph('Prenoms: ' + result[0]['pren'])

    document.add_paragraph('')

    table = document.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Matiere'
    hdr_cells[1].text = 'Coefficient'
    hdr_cells[2].text = 'Note'
    hdr_cells[3].text = 'Total'

    totalNote = 0
    coefficientNote = 0
    totals = 0
    for item in result:
        row_cells = table.add_row().cells
        row_cells[0].text = str(item['lib'])
        row_cells[1].text = str(item['coef'])
        row_cells[2].text = str(item['note'])
        row_cells[3].text = str(item['note']*item['coef'])
        totalNote += item['note']
        coefficientNote += item['coef']
        totals += (item['coef']*item['note'])

    totalMoyenne = totals / coefficientNote
    ft_cells = table.add_row().cells
    ft_cells[0].text = str("")
    ft_cells[1].text = str("")
    ft_cells[2].text = str("Total")
    ft_cells[3].text = str(totalMoyenne)

    document.add_page_break()

    document.save('/tmp/demo.docx')
    process = subprocess.Popen(['libreoffice', '', '/tmp/demo.docx'])
    time.sleep(2)
    process.kill()
    return jsonify({'item': 'bulletin ok'}), 201


if __name__ == '__main__':
    app.run(debug=True)
